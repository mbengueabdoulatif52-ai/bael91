"""
export/note_calcul.py
Génération de la note de calcul complète BAEL 91
Format texte structuré — tous les éléments avec cheminement complet
"""
import math
import io
from datetime import datetime
from pathlib import Path
import sys
sys.path.insert(0, str(Path(__file__).parent.parent))


# ── Helpers formatage ─────────────────────────────────────────────────────────
SEP1 = "═" * 70
SEP2 = "─" * 70
SEP3 = "·" * 70

def _titre1(texte):
    return f"\n{SEP1}\n{texte.center(70)}\n{SEP1}\n"

def _titre2(texte):
    return f"\n{SEP2}\n{texte}\n{SEP2}\n"

def _titre3(texte):
    return f"\n  {texte}\n  {'─'*50}\n"

def _ok(cond, msg_ok="✓ OK", msg_nok="✗ REVOIR"):
    return msg_ok if cond else msg_nok

def _statut(texte):
    if not texte: return ""
    if "REVOIR" in texte.upper(): return f"  ❌ {texte}"
    if "OK" in texte.upper(): return f"  ✅ {texte}"
    return f"  → {texte}"


# ── En-tête ───────────────────────────────────────────────────────────────────
def _entete(projet, res):
    mat = projet.materiaux
    fbu = 0.85 * mat.fc28 / mat.gammab
    fsu = mat.fe / mat.gammas
    ftj = 0.06 * mat.fc28 + 0.6
    Ec  = 11000 * mat.fc28**(1/3)

    nb_alertes = (sum(1 for r in res.poutres if r.alerte) +
                  sum(1 for r in res.poteaux if r.alerte_am))

    lignes = []
    lignes.append(_titre1("NOTE DE CALCUL — BÉTON ARMÉ BAEL 91"))
    lignes.append(f"  Projet  : {projet.nom}")
    lignes.append(f"  Date    : {datetime.now().strftime('%d/%m/%Y  %H:%M')}")
    lignes.append(f"  Niveaux : {projet.nb_niveaux}")
    lignes.append(f"  Éléments: {len(res.poutres)} poutres · "
                  f"{len(res.poteaux)} poteaux · "
                  f"{len(res.dalles)} dalles · "
                  f"{len(res.semelles)} semelles")
    lignes.append(f"  Alertes : {nb_alertes}")

    lignes.append(_titre2("1. MATÉRIAUX"))
    lignes.append(f"  Béton :")
    lignes.append(f"    fc28  = {mat.fc28:.0f} MPa")
    lignes.append(f"    γb    = {mat.gammab:.2f}")
    lignes.append(f"    fbu   = 0.85 × {mat.fc28:.0f} / {mat.gammab:.2f} = {fbu:.2f} MPa")
    lignes.append(f"    ftj   = 0.06 × {mat.fc28:.0f} + 0.6 = {ftj:.2f} MPa")
    lignes.append(f"    Ec    = 11000 × {mat.fc28:.0f}^(1/3) = {Ec:.0f} MPa")
    lignes.append(f"    σbc   = {mat.sigmaBc:.1f} MPa  (ELS)")
    lignes.append(f"  Acier :")
    lignes.append(f"    fe    = {mat.fe:.0f} MPa")
    lignes.append(f"    γs    = {mat.gammas:.2f}")
    lignes.append(f"    fsu   = {mat.fe:.0f} / {mat.gammas:.2f} = {fsu:.2f} MPa")
    lignes.append(f"  Enrobage : poutres/poteaux c = {mat.c_poutre*100:.0f}cm  "
                  f"dalles c = {mat.c_dalle*100:.0f}cm  "
                  f"fondations c = {mat.c_fond*100:.0f}cm")
    lignes.append(f"  Sol :")
    lignes.append(f"    q_adm = {mat.q_adm:.0f} kN/m²")
    lignes.append(f"    Df    = {mat.Df:.2f}m")

    return "\n".join(lignes)


# ── Dalles ────────────────────────────────────────────────────────────────────
def _note_dalle(r, projet):
    mat = projet.materiaux
    fbu = 0.85 * mat.fc28 / mat.gammab
    fsu = mat.fe / mat.gammas
    d = next((d for d in projet.dalles if d.id == r.dalle_id), None)

    lignes = []
    lignes.append(f"\n  Dalle D{r.dalle_id}  Type: {r.type_dalle}  ({r.typH})")
    lignes.append(f"  {'─'*55}")

    if d:
        lignes.append(f"    G = {d.G:.2f} kN/m²  Q = {d.Q:.2f} kN/m²")
        qu = 1.35*d.G + 1.5*d.Q
        lignes.append(f"    qu = 1.35×{d.G:.2f} + 1.5×{d.Q:.2f} = {qu:.2f} kN/m²")

    lignes.append(f"\n    ELU Flexion :")
    lignes.append(f"      Mu_x = {r.Mu_x:.2f} kN.m/m")
    if r.Mu_y > 0:
        lignes.append(f"      Mu_y = {r.Mu_y:.2f} kN.m/m")

    lignes.append(f"\n    Armatures :")
    lignes.append(f"      As nerv. = {r.As_nerv:.2f} cm²/m")
    lignes.append(f"      As rép.  = {r.As_rep:.2f}  cm²/m")

    lignes.append(f"\n    Vérifications :")
    lignes.append(_statut(r.vH))
    lignes.append(_statut(r.vELS))

    statut = "✅ OK" if not r.alerte else "❌ REVOIR"
    lignes.append(f"\n    → {statut}")
    return "\n".join(lignes)


# ── Poutres ───────────────────────────────────────────────────────────────────
def _note_poutre(r, b, mat):
    fbu = 0.85 * mat.fc28 / mat.gammab
    fsu = mat.fe / mat.gammas
    Ec  = 11000 * mat.fc28**(1/3)

    bw = b.b if b else 0.15
    h  = b.h if b else 0.30
    L  = b.longueur if b else 0
    d  = h - mat.c_poutre
    d_mm = d * 1000
    b_mm = bw * 1000

    lignes = []
    lignes.append(f"\n  {r.etiq}   Niv.{b.niveau if b else '?'}   L = {L:.2f}m")
    lignes.append(f"  {'─'*55}")
    lignes.append(f"    Section : b = {bw*100:.0f}cm   h = {h*100:.0f}cm")
    lignes.append(f"    d = h - c = {h*100:.0f} - {mat.c_poutre*100:.0f} = {d*100:.1f}cm")

    lignes.append(f"\n    FLEXION ELU")
    lignes.append(f"      Mu = {r.Mu:.2f} kN.m")
    if r.Mu > 0 and d_mm > 0:
        Mu_Nmm = r.Mu * 1e6
        mu = Mu_Nmm / (b_mm * d_mm**2 * fbu)
        lignes.append(f"      μ = Mu/(b×d²×fbu)")
        lignes.append(f"        = {r.Mu:.2f}×10⁶ / ({b_mm:.0f}×{d_mm:.0f}²×{fbu:.2f})")
        lignes.append(f"        = {mu:.4f}  {'✅ ≤ 0.392  Pivot A' if mu<=0.392 else '❌ > 0.392'}")

        if mu <= 0.392:
            alpha = (1 - math.sqrt(max(1-2*mu, 0))) / 0.8
            z_mm  = d_mm * (1 - 0.4*alpha)
            As_min1 = 0.23*b_mm*d_mm*mat.ftj/mat.fe/100
            As_min2 = 0.001*b_mm*h*1000/100
            lignes.append(f"      α = (1-√(1-2×{mu:.4f}))/0.8 = {alpha:.4f}")
            lignes.append(f"      z = d×(1-0.4α) = {d_mm:.0f}×(1-0.4×{alpha:.4f}) = {z_mm:.1f}mm")
            As_calc = Mu_Nmm / (z_mm * fsu) / 100 if z_mm > 0 else 0
            lignes.append(f"      As = Mu/(z×fsu) = {r.Mu:.2f}×10⁶/({z_mm:.1f}×{fsu:.2f})/100")
            lignes.append(f"         = {As_calc:.2f} cm²")
            lignes.append(f"      As_min = max({As_min1:.2f}, {As_min2:.2f}) = {max(As_min1,As_min2):.2f} cm²")
            lignes.append(f"      As long. retenu = {r.As_long:.2f} cm²")

    if r.As_chap > 0:
        lignes.append(f"      As chap. = {r.As_chap:.2f} cm²  (moment sur appui)")

    lignes.append(f"\n    CISAILLEMENT ELU")
    lignes.append(f"      Tu = {r.Tu:.2f} kN")
    if r.Tu > 0 and b_mm > 0 and d_mm > 0:
        tau_u   = r.Tu * 1000 / (b_mm * d_mm)
        tau_lim = min(0.2 * mat.fc28 / mat.gammab, 5.0)
        lignes.append(f"      τu = Tu/(b×d) = {r.Tu*1000:.0f}/({b_mm:.0f}×{d_mm:.0f})")
        lignes.append(f"         = {tau_u:.3f} MPa  ≤ τ_lim={tau_lim:.2f} MPa  "
                      f"{'✅' if tau_u<=tau_lim else '❌'}")
        lignes.append(f"      At/st = {r.At_st:.2f} cm²/m")
    lignes.append(_statut(r.vCis))

    lignes.append(f"\n    FLÈCHE ELS (indicative)")
    lignes.append(_statut(r.vFleche if r.vFleche else "Non calculée"))
    lignes.append(_statut(r.vH))

    statut = "✅ SECTION OK" if not r.alerte else "❌ SECTION À REVOIR"
    lignes.append(f"\n    → {statut}")
    return "\n".join(lignes)


# ── Poteaux ───────────────────────────────────────────────────────────────────
def _note_poteau(r, b, mat):
    fbu = 0.85 * mat.fc28 / mat.gammab
    fsu = mat.fe / mat.gammas

    bw = b.b if b else 0.15
    h  = b.h if b else 0.15
    L  = b.longueur if b else 0
    lf = 0.7 * L

    lignes = []
    lignes.append(f"\n  {r.etiq}   Niv.{b.niveau if b else '?'}   H = {L:.2f}m")
    lignes.append(f"  {'─'*55}")
    lignes.append(f"    Section : b = {bw*100:.0f}cm   h = {h*100:.0f}cm")

    lignes.append(f"\n    DESCENTE DE CHARGES")
    lignes.append(f"      Nu = {r.Nu:.1f} kN  (ELU)")

    lignes.append(f"\n    FLAMBEMENT")
    lignes.append(f"      lf = 0.7 × H = 0.7 × {L:.2f} = {lf:.2f}m")
    i  = min(bw, h) / math.sqrt(12)
    lam = lf / i
    lignes.append(f"      i = min(b,h)/√12 = {min(bw,h)*100:.0f}/√12 = {i*100:.2f}cm")
    lignes.append(f"      λ = lf/i = {lf*100:.0f}/{i*100:.2f} = {lam:.1f}")
    lignes.append(f"      λ = {lam:.1f} {'✅ ≤ 70' if lam<=70 else '❌ > 70'}  "
                  f"→ {'Méthode forfaitaire' if lam<=70 else 'Hors domaine'}")

    if lam <= 70:
        alpha = 0.85 / (1 + 0.2*(lam/35)**2)
        lignes.append(f"      α = 0.85/(1+0.2×(λ/35)²) = {alpha:.4f}")

        N_beton = alpha * fbu * bw*1000 * h*1000 / 1000
        lignes.append(f"\n    DIMENSIONNEMENT")
        lignes.append(f"      N_béton = α×fbu×b×h = {alpha:.4f}×{fbu:.2f}×"
                      f"{bw*1000:.0f}×{h*1000:.0f}/1000 = {N_beton:.1f} kN")
        As_max = 0.05 * bw*1000 * h*1000 / 100
        As_min = max(0.002 * bw*1000 * h*1000 / 100, 4*0.503)
        lignes.append(f"      As = (Nu - N_béton)/fsu")
        As_calc = max((r.Nu*1000 - N_beton*1000) / fsu / 100, 0)
        lignes.append(f"         = ({r.Nu*1000:.0f} - {N_beton*1000:.0f})/{fsu:.2f}/100")
        lignes.append(f"         = {As_calc:.2f} cm²")
        lignes.append(f"      As_min = {As_min:.2f} cm²")
        lignes.append(f"      As_max = 0.05×{bw*100:.0f}×{h*100:.0f}/100 = {As_max:.2f} cm²")
        lignes.append(f"      As retenu = {r.As:.2f} cm²  "
                      f"{'✅' if r.As<=As_max else '❌ > As_max'}")

    lignes.append(f"\n    Vérification ELU : {r.vL}")
    lignes.append(f"    {r.vS}")

    statut = "✅ OK" if not r.alerte_am else "❌ REVOIR"
    lignes.append(f"\n    → {statut}")
    return "\n".join(lignes)


# ── Semelles ──────────────────────────────────────────────────────────────────
def _note_semelle(s, pot, mat, noms_pots):
    fbu = 0.85 * mat.fc28 / mat.gammab
    fsu = mat.fe / mat.gammas
    q_adm = s.q_adm_loc if s.q_adm_loc > 0 else mat.q_adm

    b_pot = pot.b if pot else 0.15
    h_pot = pot.h if pot else 0.15
    nom   = noms_pots.get(s.id_poteau, f"C{s.id_poteau}")

    lignes = []
    lignes.append(f"\n  Semelle sous {nom}   "
                  f"{'Excentrique' if s.ex!=0 or s.ey!=0 else 'Centrée'}")
    lignes.append(f"  {'─'*55}")
    lignes.append(f"    Poteau : b = {b_pot*100:.0f}cm   h = {h_pot*100:.0f}cm")
    lignes.append(f"    Nu_ELU = {s.Nu_ELU:.1f} kN")
    lignes.append(f"    Nu_ser = {s.Nu_ELU:.1f}/1.35 = {s.Nu_ser:.1f} kN")
    lignes.append(f"    q_adm  = {q_adm:.0f} kN/m²")

    if s.ex != 0 or s.ey != 0:
        lignes.append(f"    ex = {s.ex}  ey = {s.ey}  (convention -1/0/+1)")
        lignes.append(f"    ex_reel = {getattr(s,'ex_reel',s.ex):.3f}m  "
                      f"ey_reel = {getattr(s,'ey_reel',s.ey):.3f}m")

    lignes.append(f"\n    DIMENSIONS EN PLAN")
    lignes.append(f"      B = √(Nu_ser × 1.10 / q_adm)  (point départ)")
    lignes.append(f"        = √({s.Nu_ser:.1f} × 1.10 / {q_adm:.0f})")
    lignes.append(f"        → B = L = {s.B:.2f}m  (arrondi au 5cm)")

    lignes.append(f"\n    VÉRIFICATION SOL")
    lignes.append(f"      q_max = {s.q_max:.1f} kN/m²  "
                  f"{'✅ ≤ q_adm' if s.q_max<=q_adm else '❌ > q_adm'}")
    lignes.append(f"      q_min = {s.q_min:.1f} kN/m²  "
                  f"{'✅' if s.q_min>=0 else '❌ SOULÈVEMENT'}")

    lignes.append(f"\n    HAUTEUR DE LA SEMELLE")
    debord = (s.B - b_pot) / 2
    tau_lim = 0.07 * mat.fc28 / mat.gammab
    V = s.q_max * debord
    d_u_min = V / (tau_lim * 1000)
    lignes.append(f"      Débord = (B - b_pot)/2 = ({s.B:.2f} - {b_pot:.3f})/2 = {debord:.3f}m")
    lignes.append(f"      τ_lim  = 0.07×fc28/γb = {tau_lim:.3f} MPa")
    lignes.append(f"      V      = q_max × débord = {s.q_max:.1f} × {debord:.3f} = {V:.2f} kN/m")
    lignes.append(f"      d_u    = V/(τ_lim×1000) = {d_u_min*100:.1f}cm")
    lignes.append(f"      e_sem  = {s.e_sem*100:.0f}cm  (retenu)")

    lignes.append(f"\n    FERRAILLAGE")
    d_u = s.e_sem - mat.c_fond
    Mu_s = s.q_max * debord**2 / 2
    lignes.append(f"      d    = {s.e_sem*100:.0f} - {mat.c_fond*100:.0f} = {d_u*100:.1f}cm")
    lignes.append(f"      Mu   = q_max × débord²/2 = {s.q_max:.1f}×{debord:.3f}²/2 = {Mu_s:.4f} kN.m/ml")
    lignes.append(f"      As   = Mu×10⁶/(0.9×d×fsu)/100")
    As_c = Mu_s * 1e6 / (0.9 * d_u*1000 * fsu) / 100
    lignes.append(f"         = {Mu_s:.4f}×10⁶/(0.9×{d_u*1000:.0f}×{fsu:.2f})/100")
    lignes.append(f"         = {As_c:.3f} cm²/ml")
    lignes.append(f"      Asx retenu = {s.Asx:.2f} cm²/ml")

    lignes.append(f"\n    AMORCES")
    lignes.append(f"      {s.nb_amorce}HA{s.phi_amorce}  "
                  f"ls = 40×φ = {s.ls_amorce*100:.0f}cm")

    # Longrines
    for direc, b_l, h_l, As_l, Mu_l, vM, vers in [
        ("X", s.b_long_X, s.h_long_X, s.long_X_As,
         s.long_X_Mu, s.long_X_vM, s.long_X_vers),
        ("Y", s.b_long_Y, s.h_long_Y, s.long_Y_As,
         s.long_Y_Mu, s.long_Y_vM, s.long_Y_vers),
    ]:
        e_dir = getattr(s,'ex_reel',s.ex) if direc=="X" \
                else getattr(s,'ey_reel',s.ey)
        if abs(s.ex if direc=="X" else s.ey) == 0: continue
        nom_dest = noms_pots.get(vers, f"C{vers}")
        lignes.append(f"\n    LONGRINE {direc} → {nom_dest}")
        lignes.append(f"      Section : {b_l*100:.0f}×{h_l*100:.0f}cm")
        lignes.append(f"      |e_reel| = {abs(e_dir):.3f}m")
        lignes.append(f"      M = Nu_ser × e = {s.Nu_ser:.1f}×{abs(e_dir):.3f} = {Mu_l:.1f} kN.m")
        if b_l > 0 and h_l > 0:
            d_l = h_l - mat.c_poutre
            mu_l = Mu_l / (b_l*1000 * (d_l*1000)**2 * fbu/1e6)
            lignes.append(f"      d = {h_l*100:.0f}-{mat.c_poutre*100:.0f} = {d_l*100:.1f}cm")
            lignes.append(f"      μ = {mu_l:.4f}  "
                          f"{'✅ ≤ 0.392' if mu_l<=0.392 else '❌ > 0.392'}")
            lignes.append(f"      As long. = {As_l:.2f} cm²  "
                          f"As chap. = {As_l*0.5:.2f} cm²")

    alertes = getattr(s, 'alertes', [])
    statut = "✅ OK" if not alertes else "\n    ".join(alertes)
    lignes.append(f"\n    → {statut}")
    return "\n".join(lignes)


# ── Récapitulatif alertes ─────────────────────────────────────────────────────
def _recapitulatif(res, projet):
    lignes = []
    lignes.append(_titre2("6. RÉCAPITULATIF DES ALERTES"))

    # Poutres
    poutres_alerte = [r for r in res.poutres if r.alerte]
    lignes.append(f"\n  POUTRES : {len(poutres_alerte)} alerte(s)")
    for r in poutres_alerte:
        al = []
        if r.mu_r > 0.392: al.append(f"μ={r.mu_r:.3f}>0.392")
        if r.vCis and "REVOIR" in r.vCis: al.append("Cisaillement")
        if r.vFleche and "REVOIR" in r.vFleche: al.append("Flèche")
        if r.vH and "REVOIR" in r.vH: al.append(r.vH)
        lignes.append(f"    {r.etiq:35s} ❌ {' | '.join(al) if al else 'REVOIR'}")

    # Poteaux
    poteaux_alerte = [r for r in res.poteaux if r.alerte_am]
    lignes.append(f"\n  POTEAUX : {len(poteaux_alerte)} alerte(s)")
    for r in poteaux_alerte:
        b_ref = next((b for b in projet.barres if b.id==r.barre_id), None)
        As_max = (0.05*b_ref.b*1000*b_ref.h*1000/100 if b_ref else 999)
        al = []
        if r.As > As_max: al.append(f"As={r.As:.2f}>As_max={As_max:.2f}cm²")
        if r.lam > 70: al.append(f"λ={r.lam:.0f}>70")
        if not al: al.append("Recouvrement amorces")
        lignes.append(f"    {r.etiq:35s} ❌ {' | '.join(al)}")

    # Semelles
    sem_alerte = [s for s in res.semelles if getattr(s,'alertes',[])]
    noms = {b.id: b.nom for b in projet.barres if b.type_elem=="poteau"}
    lignes.append(f"\n  SEMELLES : {len(sem_alerte)} alerte(s)")
    for s in sem_alerte:
        nom = noms.get(s.id_poteau, f"C{s.id_poteau}")
        for a in getattr(s,'alertes',[]):
            lignes.append(f"    {nom:30s} {a}")

    total = len(poutres_alerte) + len(poteaux_alerte) + len(sem_alerte)
    lignes.append(f"\n  TOTAL : {total} alerte(s)")
    return "\n".join(lignes)


# ── Fonction principale ───────────────────────────────────────────────────────
def generer_note_calcul(res, projet) -> str:
    """Génère la note de calcul complète en texte."""
    mat    = projet.materiaux
    noms   = {b.id: b.nom for b in projet.barres if b.type_elem=="poteau"}
    b_map  = {b.id: b for b in projet.barres}

    lignes = []

    # En-tête + matériaux
    lignes.append(_entete(projet, res))

    # ── Dalles ────────────────────────────────────────────────────────────────
    lignes.append(_titre2("2. DALLES"))
    lignes.append(f"  Nombre de dalles calculées : {len(res.dalles)}\n")
    for r in res.dalles:
        lignes.append(_note_dalle(r, projet))

    # ── Poutres par niveau ────────────────────────────────────────────────────
    lignes.append(_titre2("3. POUTRES"))
    niveaux = sorted(set(
        b.niveau for b in projet.barres if b.type_elem=="poutre"
    ), reverse=True)

    for niv in niveaux:
        lignes.append(f"\n  ── Niveau {niv} ──")
        poutres_niv = [r for r in res.poutres
                       if b_map.get(r.barre_id,
                          type('',(),{'niveau':0})()).niveau == niv]
        for r in poutres_niv:
            b = b_map.get(r.barre_id)
            lignes.append(_note_poutre(r, b, mat))

    # ── Poteaux par niveau ────────────────────────────────────────────────────
    lignes.append(_titre2("4. POTEAUX"))
    niveaux_p = sorted(set(
        b.niveau for b in projet.barres if b.type_elem=="poteau"
    ), reverse=True)

    for niv in niveaux_p:
        lignes.append(f"\n  ── Niveau {niv} ──")
        poteaux_niv = [r for r in res.poteaux
                       if b_map.get(r.barre_id,
                          type('',(),{'niveau':0})()).niveau == niv]
        for r in poteaux_niv:
            b = b_map.get(r.barre_id)
            lignes.append(_note_poteau(r, b, mat))

    # ── Fondations ────────────────────────────────────────────────────────────
    lignes.append(_titre2("5. FONDATIONS"))
    pots_n1 = {b.id: b for b in projet.barres
               if b.type_elem=="poteau" and b.niveau==1}
    for s in res.semelles:
        pot = pots_n1.get(s.id_poteau)
        lignes.append(_note_semelle(s, pot, mat, noms))

    # ── Récapitulatif ─────────────────────────────────────────────────────────
    lignes.append(_recapitulatif(res, projet))

    lignes.append(f"\n{SEP1}")
    lignes.append("FIN DE LA NOTE DE CALCUL".center(70))
    lignes.append(f"{SEP1}\n")

    return "\n".join(lignes)


# ── Export Word (.docx) ───────────────────────────────────────────────────────
def generer_note_docx(res, projet) -> bytes:
    """Génère la note de calcul en format Word (.docx)."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx requis : pip install python-docx")

    note_txt = generer_note_calcul(res, projet)
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Courier New'
    style_normal.font.size = Pt(8)

    # Titre principal
    titre = doc.add_heading('NOTE DE CALCUL — BÉTON ARMÉ BAEL 91', level=1)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titre.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    # Info projet
    mat = projet.materiaux
    doc.add_paragraph(
        f"Projet : {projet.nom}  |  "
        f"Date : {__import__('datetime').datetime.now().strftime('%d/%m/%Y')}"
        f"  |  Niveaux : {projet.nb_niveaux}"
    )
    doc.add_paragraph()

    # Parser le texte ligne par ligne
    for line in note_txt.split('\n'):
        stripped = line.strip()

        # Titres de section (══ ou ──)
        if stripped.startswith('══') or stripped.startswith('──'):
            continue
        if stripped and all(c in '═─·' for c in stripped):
            continue

        # Titres numérotés (1. 2. ...)
        if len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == '.':
            p = doc.add_heading(stripped, level=2)
            p.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
            continue

        # Alertes
        if '❌' in stripped or '✅' in stripped or '⚠' in stripped:
            p = doc.add_paragraph(line)
            run = p.runs[0] if p.runs else p.add_run(line)
            if '❌' in stripped:
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                run.bold = True
            elif '✅' in stripped:
                run.font.color.rgb = RGBColor(0x37, 0x5C, 0x23)
            elif '⚠' in stripped:
                run.font.color.rgb = RGBColor(0xBF, 0x8F, 0x00)
            continue

        # Lignes d'éléments (Pt ou Poteau)
        if (stripped.startswith('Pt') or stripped.startswith('P1_') or
                stripped.startswith('Semelle')):
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].bold = True
            continue

        # Ligne normale
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Export PDF ────────────────────────────────────────────────────────────────
def generer_note_pdf(res, projet) -> bytes:
    """Génère la note de calcul en format PDF."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Paragraph,
                                        Spacer, HRFlowable)
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
    except ImportError:
        raise ImportError("reportlab requis : pip install reportlab")

    note_txt = generer_note_calcul(res, projet)
    buf = io.BytesIO()

    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    style_mono = ParagraphStyle(
        'Mono', fontName='Courier', fontSize=7,
        leading=9, leftIndent=0)
    style_titre = ParagraphStyle(
        'Titre', fontName='Helvetica-Bold', fontSize=14,
        alignment=TA_CENTER, textColor=colors.HexColor('#1F3864'),
        spaceAfter=6)
    style_section = ParagraphStyle(
        'Section', fontName='Helvetica-Bold', fontSize=9,
        textColor=colors.HexColor('#1F3864'), spaceBefore=8, spaceAfter=4)
    style_ok = ParagraphStyle(
        'OK', fontName='Courier-Bold', fontSize=7,
        textColor=colors.HexColor('#375C23'), leading=9)
    style_err = ParagraphStyle(
        'Err', fontName='Courier-Bold', fontSize=7,
        textColor=colors.HexColor('#C00000'), leading=9)
    style_warn = ParagraphStyle(
        'Warn', fontName='Courier-Bold', fontSize=7,
        textColor=colors.HexColor('#BF8F00'), leading=9)

    story = []
    story.append(Paragraph('NOTE DE CALCUL — BÉTON ARMÉ BAEL 91', style_titre))
    mat = projet.materiaux
    from datetime import datetime as _dt
    story.append(Paragraph(
        f"Projet : {projet.nom}  |  "
        f"Date : {_dt.now().strftime('%d/%m/%Y')}  |  "
        f"Niveaux : {projet.nb_niveaux}",
        styles['Normal']))
    story.append(Spacer(1, 0.3*cm))

    for line in note_txt.split('\n'):
        stripped = line.strip()

        # Séparateurs
        if all(c in '═─·' for c in stripped) and stripped:
            story.append(HRFlowable(
                width="100%", thickness=0.5,
                color=colors.HexColor('#CCCCCC'),
                spaceAfter=2, spaceBefore=2))
            continue

        # Titres numérotés
        if len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == '.':
            story.append(Paragraph(stripped, style_section))
            continue

        # Lignes vides
        if not stripped:
            story.append(Spacer(1, 0.15*cm))
            continue

        # Alertes colorées
        if '❌' in stripped:
            safe = stripped.replace('❌', '[ERREUR]').replace('<','&lt;').replace('>','&gt;')
            story.append(Paragraph(safe, style_err))
        elif '✅' in stripped:
            safe = stripped.replace('✅', '[OK]').replace('<','&lt;').replace('>','&gt;')
            story.append(Paragraph(safe, style_ok))
        elif '⚠' in stripped:
            safe = stripped.replace('⚠', '[AVERT.]').replace('<','&lt;').replace('>','&gt;')
            story.append(Paragraph(safe, style_warn))
        else:
            safe = line.replace('<','&lt;').replace('>','&gt;')
            story.append(Paragraph(safe, style_mono))

    doc.build(story)
    buf.seek(0)
    return buf.read()
